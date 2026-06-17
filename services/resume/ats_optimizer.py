"""
Resume Generation & ATS Optimization Engine
Creates tailored resumes in DOCX format optimized for ATS parsing
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Dict, List, Optional
import re
from datetime import datetime, timezone
import json

from utils.logger import logger


class ATSResumeEngine:
    """
    ATS-optimized resume generator
    Creates clean, parseable DOCX resumes
    """
    
    # ATS-friendly fonts
    SAFE_FONTS = ['Calibri', 'Arial', 'Times New Roman', 'Georgia']
    
    # Standard section headings for ATS
    STANDARD_SECTIONS = {
        'summary': 'PROFESSIONAL SUMMARY',
        'experience': 'PROFESSIONAL EXPERIENCE',
        'education': 'EDUCATION',
        'skills': 'TECHNICAL SKILLS',
        'projects': 'PROJECTS',
        'certifications': 'CERTIFICATIONS'
    }
    
    def __init__(self, base_resume_data: Dict):
        """
        Initialize with base resume data
        
        Args:
            base_resume_data: Parsed resume data structure
        """
        self.base_data = base_resume_data
        
    def generate_tailored_resume(self,
                                job_description: str,
                                job_title: str,
                                keywords: List[str],
                                variant: str = 'balanced') -> tuple[Document, Dict]:
        """
        Generate ATS-optimized resume tailored to job
        
        Args:
            job_description: Full job description text
            job_title: Target job title
            keywords: Keywords to emphasize
            variant: 'conservative', 'balanced', or 'aggressive'
        
        Returns:
            (Document object, metadata dict)
        """
        logger.info(f"📝 Generating {variant} resume variant")
        
        # Create new document
        doc = Document()
        
        # Set margins (1 inch all around - ATS standard)
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Optimize data for this job
        optimized_data = self._optimize_content(
            self.base_data,
            job_description,
            job_title,
            keywords,
            variant
        )
        
        # Add sections
        self._add_header(doc, optimized_data)
        self._add_summary(doc, optimized_data, job_title, keywords)
        self._add_skills(doc, optimized_data, keywords)
        self._add_experience(doc, optimized_data, keywords)
        self._add_projects(doc, optimized_data, keywords)
        self._add_education(doc, optimized_data)
        
        # Calculate ATS score
        ats_score = self._calculate_ats_score(doc, job_description, keywords)
        
        metadata = {
            'variant': variant,
            'keywords_added': optimized_data.get('keywords_added', []),
            'ats_score': ats_score,
            'generated_at': datetime.now(timezone.utc).isoformat(),
            'job_title': job_title
        }
        
        logger.info(f"✓ Resume generated | ATS Score: {ats_score}")
        
        return doc, metadata
    
    def _optimize_content(self, base_data: Dict, job_desc: str,
                         job_title: str, keywords: List[str],
                         variant: str) -> Dict:
        """Optimize resume content based on job requirements"""
        
        optimized = base_data.copy()
        keywords_added = []
        
        # Adjust professional summary
        if variant in ['balanced', 'aggressive']:
            summary = optimized.get('summary', '')
            
            # Add job title if not present
            if job_title.lower() not in summary.lower():
                summary = f"{job_title} " + summary
                keywords_added.append(f"Added '{job_title}' to summary")
            
            # Inject relevant keywords naturally
            for keyword in keywords[:5]:  # Top 5 keywords
                if keyword.lower() not in summary.lower():
                    # Find natural place to add
                    if 'experience' in summary.lower():
                        summary = summary.replace(
                            'experience',
                            f'experience in {keyword}'
                        )
                        keywords_added.append(f"Added '{keyword}' to summary")
                        break
            
            optimized['summary'] = summary
        
        # Reorder skills to match job requirements
        if 'skills' in optimized and keywords:
            skills = optimized['skills']
            primary = skills.get('primary', [])
            
            # Move matching skills to front
            matched_skills = [s for s in primary if any(k.lower() in s.lower() for k in keywords)]
            other_skills = [s for s in primary if s not in matched_skills]
            
            skills['primary'] = matched_skills + other_skills
            optimized['skills'] = skills
        
        # Enhance experience descriptions
        if variant == 'aggressive':
            experiences = optimized.get('experience', [])
            for exp in experiences:
                achievements = exp.get('achievements', [])
                for i, achievement in enumerate(achievements):
                    # Add relevant keywords to achievements
                    for keyword in keywords[:3]:
                        if keyword.lower() not in achievement.lower():
                            # Add keyword naturally
                            achievement = achievement.replace(
                                'using',
                                f'using {keyword} and'
                            )
                            keywords_added.append(f"Enhanced achievement with '{keyword}'")
                            break
                    achievements[i] = achievement
                exp['achievements'] = achievements
        
        optimized['keywords_added'] = keywords_added
        return optimized
    
    def _add_header(self, doc: Document, data: Dict):
        """Add header with contact info"""
        # Name
        name_para = doc.add_paragraph()
        name_run = name_para.add_run(data.get('name', 'Unknown'))
        name_run.font.size = Pt(16)
        name_run.font.bold = True
        name_run.font.name = 'Calibri'
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Contact info - single line
        contact_parts = []
        if data.get('email'):
            contact_parts.append(data['email'])
        if data.get('phone'):
            contact_parts.append(data['phone'])
        if data.get('location'):
            contact_parts.append(data['location'])
        if data.get('linkedin'):
            contact_parts.append(data['linkedin'])
        if data.get('github'):
            contact_parts.append(data['github'])
        
        contact_para = doc.add_paragraph(' • '.join(contact_parts))
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._set_font(contact_para, size=10)
        
        # Add spacing
        doc.add_paragraph()
    
    def _add_summary(self, doc: Document, data: Dict, 
                    job_title: str, keywords: List[str]):
        """Add professional summary"""
        summary_text = data.get('summary', '')
        
        if not summary_text and data.get('experience'):
            # Generate from experience if not provided
            summary_text = f"{job_title} with expertise in {', '.join(keywords[:5])}."
        
        if summary_text:
            self._add_section_heading(doc, self.STANDARD_SECTIONS['summary'])
            para = doc.add_paragraph(summary_text)
            self._set_font(para)
            doc.add_paragraph()
    
    def _add_skills(self, doc: Document, data: Dict, keywords: List[str]):
        """Add skills section with categorization"""
        skills = data.get('skills', {})
        
        if not skills:
            return
        
        self._add_section_heading(doc, self.STANDARD_SECTIONS['skills'])
        
        # Primary skills
        primary = skills.get('primary', [])
        if primary:
            para = doc.add_paragraph()
            run = para.add_run('Primary: ')
            run.bold = True
            para.add_run(' • '.join(primary))
            self._set_font(para)
        
        # Secondary skills
        secondary = skills.get('secondary', [])
        if secondary:
            para = doc.add_paragraph()
            run = para.add_run('Additional: ')
            run.bold = True
            para.add_run(' • '.join(secondary))
            self._set_font(para)
        
        doc.add_paragraph()
    
    def _add_experience(self, doc: Document, data: Dict, keywords: List[str]):
        """Add work experience"""
        experiences = data.get('experience', [])
        
        if not experiences:
            return
        
        self._add_section_heading(doc, self.STANDARD_SECTIONS['experience'])
        
        for exp in experiences:
            # Job title and company
            title_para = doc.add_paragraph()
            title_run = title_para.add_run(f"{exp.get('title', 'Unknown')} | {exp.get('company', 'Unknown')}")
            title_run.bold = True
            self._set_font(title_para, size=11)
            
            # Date and location
            date_para = doc.add_paragraph()
            date_text = f"{exp.get('start_date', 'N/A')} – {exp.get('end_date', 'Present')}"
            if exp.get('location'):
                date_text += f" • {exp['location']}"
            date_para.add_run(date_text)
            date_para.style = 'List Bullet'
            self._set_font(date_para, size=10, italic=True)
            
            # Achievements
            for achievement in exp.get('achievements', []):
                para = doc.add_paragraph(achievement, style='List Bullet')
                self._set_font(para)
            
            doc.add_paragraph()  # Spacing between jobs
    
    def _add_projects(self, doc: Document, data: Dict, keywords: List[str]):
        """Add projects section"""
        projects = data.get('projects', [])
        
        if not projects:
            return
        
        self._add_section_heading(doc, self.STANDARD_SECTIONS['projects'])
        
        for proj in projects:
            # Project title
            title_para = doc.add_paragraph()
            title_run = title_para.add_run(proj.get('name', 'Unknown Project'))
            title_run.bold = True
            
            # Tech stack
            if proj.get('technologies'):
                title_para.add_run(f" | {proj['technologies']}")
            
            self._set_font(title_para, size=11)
            
            # Description
            if proj.get('description'):
                para = doc.add_paragraph(proj['description'], style='List Bullet')
                self._set_font(para)
        
        doc.add_paragraph()
    
    def _add_education(self, doc: Document, data: Dict):
        """Add education section"""
        education = data.get('education', [])
        
        if not education:
            return
        
        self._add_section_heading(doc, self.STANDARD_SECTIONS['education'])
        
        for edu in education:
            # Degree and institution
            para = doc.add_paragraph()
            degree_run = para.add_run(f"{edu.get('degree', 'Unknown')} | {edu.get('institution', 'Unknown')}")
            degree_run.bold = True
            self._set_font(para, size=11)
            
            # Date and GPA
            details = []
            if edu.get('graduation_date'):
                details.append(edu['graduation_date'])
            if edu.get('gpa'):
                details.append(f"GPA: {edu['gpa']}")
            
            if details:
                detail_para = doc.add_paragraph(' • '.join(details))
                self._set_font(detail_para, size=10, italic=True)
    
    def _add_section_heading(self, doc: Document, heading: str):
        """Add section heading in ATS-friendly format"""
        para = doc.add_paragraph(heading)
        run = para.runs[0]
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.name = 'Calibri'
        run.font.color.rgb = RGBColor(0, 0, 0)
        
        # Add subtle underline (ATS safe)
        para.add_run('_' * 80)
        
    def _set_font(self, paragraph, size: int = 11, bold: bool = False, italic: bool = False):
        """Set font properties for ATS compatibility"""
        for run in paragraph.runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(size)
            run.font.bold = bold
            run.font.italic = italic
    
    def _calculate_ats_score(self, doc: Document, job_desc: str, 
                            keywords: List[str]) -> float:
        """
        Calculate ATS compatibility score (0-100)
        Based on keyword matching and format compliance
        """
        score = 100.0
        
        # Extract all text from document
        doc_text = '\n'.join([para.text for para in doc.paragraphs]).lower()
        job_desc_lower = job_desc.lower()
        
        # Keyword matching (60% weight)
        keyword_score = 0
        for keyword in keywords[:20]:  # Top 20 keywords
            if keyword.lower() in doc_text:
                keyword_score += 3  # 3 points per keyword
        
        keyword_score = min(keyword_score, 60)  # Cap at 60
        
        # Format compliance (40% weight)
        format_score = 40
        
        # Check for problematic elements
        if len(doc.tables) > 0:
            format_score -= 10  # Tables can break ATS
        
        # Check for standard section headings
        required_sections = ['experience', 'education', 'skills']
        for section in required_sections:
            if section not in doc_text:
                format_score -= 5
        
        # Check for dates in experience
        date_pattern = r'\d{4}'
        if not re.search(date_pattern, doc_text):
            format_score -= 5
        
        # Check for bullet points (good)
        if '•' in doc_text or doc_text.count('\n') > 20:
            format_score += 5
        
        total_score = keyword_score + format_score
        return min(total_score, 100)
    
    def save_resume(self, doc: Document, filepath: str):
        """Save resume to file"""
        doc.save(filepath)
        logger.info(f"✓ Resume saved to {filepath}")


def parse_satyam_resume() -> Dict:
    """Parse Satyam's resume into structured data"""
    return {
        'name': 'Satyam Shivam',
        'email': 'shivamsatyam35@gmail.com',
        'phone': '+91 9852015381',
        'location': 'New Delhi, India',
        'linkedin': 'linkedin.com/in/satyam',
        'github': 'github.com/satyam',
        'portfolio': 'portfolio-url',
        
        'summary': 'AI Engineer specializing in production RAG systems, LangChain pipelines, and FastAPI backends. Built scalable GenAI applications processing 1M+ tokens with 89% retrieval accuracy.',
        
        'skills': {
            'primary': ['LangChain', 'RAG', 'OpenAI API', 'CrewAI', 'FastAPI', 'Python', 'Docker'],
            'secondary': ['AWS', 'PostgreSQL', 'FAISS', 'ChromaDB', 'React', 'Node.js']
        },
        
        'experience': [
            {
                'title': 'AI Developer Intern',
                'company': 'Asvix',
                'location': 'Remote',
                'start_date': 'Jan 2026',
                'end_date': 'Apr 2026',
                'achievements': [
                    'Engineered embedding pipeline for DigiLab AI chatbot, optimizing FAISS vector retrieval to improve medical query response relevance by 23%',
                    'Built context-aware response modules using LangChain, reducing hallucination rate from 18% to 11%; supported 500+ daily queries at 99.2% uptime'
                ]
            },
            {
                'title': 'AI Chatbot Development Intern',
                'company': 'Cloudily Scripts',
                'location': 'Remote',
                'start_date': 'Jun 2025',
                'end_date': 'Jul 2025',
                'achievements': [
                    'Architected production RAG pipeline processing 100+ page PDFs with 91% accuracy, cutting support tickets by 35% through FAISS IVF128 indexing',
                    'Optimized query latency by 79% (8.2s to 1.7s) via caching and parallel embedding; Dockerized services, reducing image size by 60%'
                ]
            },
            {
                'title': 'Cloud Engineering Intern',
                'company': 'IPtechhub',
                'location': 'Remote',
                'start_date': 'May 2024',
                'end_date': 'Jul 2024',
                'achievements': [
                    'Deployed containerized ML services on AWS EC2 with auto-scaling, handling 500+ daily requests at 99.5% uptime',
                    'Automated CI/CD via GitHub Actions, cutting deployment time by 87.5% from 2 hours to 15 minutes'
                ]
            }
        ],
        
        'projects': [
            {
                'name': 'Production RAG Pipeline',
                'technologies': 'LangChain, FAISS, FastAPI, ChromaDB, Docker',
                'description': 'Engineered scalable RAG system processing 1M+ tokens with hybrid retrieval, achieving 89% relevance with 850ms latency and 6% hallucination rate'
            },
            {
                'name': 'Multi-Agent Marketing Campaign Creator',
                'technologies': 'Python, CrewAI, Groq API, Pydantic',
                'description': 'Built autonomous 4-agent system generating full campaigns in 3-5 minutes versus 2-3 human days, reducing cost by 94%'
            },
            {
                'name': 'HybridAI Syntax Error Detection',
                'technologies': 'Python AST Parser, OpenAI GPT-4',
                'description': 'Designed dual-mode detection engine combining AST parsing with GPT-4 fallback, reducing debugging time by 40%'
            }
        ],
        
        'education': [
            {
                'degree': 'B.Tech in Computer Science & Engineering - AI & ML Specialization',
                'institution': 'United Institute of Technology',
                'graduation_date': '2026',
                'gpa': None
            }
        ]
    }


if __name__ == "__main__":
    # Test resume generation
    base_data = parse_satyam_resume()
    engine = ATSResumeEngine(base_data)
    
    # Generate test resume
    doc, metadata = engine.generate_tailored_resume(
        job_description="Looking for AI Engineer with LangChain and RAG experience...",
        job_title="Senior AI Engineer",
        keywords=['LangChain', 'RAG', 'FastAPI', 'Python', 'AWS'],
        variant='balanced'
    )
    
    engine.save_resume(doc, '/tmp/test_resume.docx')
    print(f"✓ Test resume generated | ATS Score: {metadata['ats_score']}")
